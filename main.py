from fastapi import FastAPI, Depends, HTTPException, Request, Form
from fastapi.responses import HTMLResponse, RedirectResponse, FileResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from sqlalchemy.orm import Session
from sqlalchemy import func
from datetime import datetime
from typing import Optional
from docx import Document
from docx.shared import Pt, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
import os
import tempfile

from src.database import engine, get_db, Base
from src.models import (
    Department, Employee, Cartridge, Warehouse, Box, 
    CartridgeLocation, ServiceNote, CartridgeMovement
)
from src.schemas import (
    DepartmentCreate, EmployeeCreate, CartridgeCreate, 
    WarehouseCreate, BoxCreate, ServiceNoteCreate,
    CartridgeLocationCreate
)

Base.metadata.create_all(bind=engine)

app = FastAPI(title="Учет картриджей")

templates = Jinja2Templates(directory="templates")


def generate_note_number(db: Session) -> str:
    year = datetime.now().year
    count = db.query(ServiceNote).filter(
        func.extract('year', ServiceNote.created_date) == year
    ).count()
    return f"КАРТ-{year}-{str(count + 1).zfill(3)}"


@app.get("/", response_class=HTMLResponse)
async def dashboard(request: Request, db: Session = Depends(get_db)):
    total_cartridges = db.query(func.sum(CartridgeLocation.quantity)).filter(
        CartridgeLocation.status == "на складе"
    ).scalar() or 0
    
    in_use = db.query(func.sum(CartridgeLocation.quantity)).filter(
        CartridgeLocation.status == "выдано"
    ).scalar() or 0
    
    recent_notes = db.query(ServiceNote).order_by(
        ServiceNote.created_date.desc()
    ).limit(5).all()
    
    low_stock_boxes = db.query(Box).filter(Box.current_count < 3).all()
    
    return templates.TemplateResponse("dashboard.html", {
        "request": request,
        "total_cartridges": total_cartridges,
        "in_use": in_use,
        "recent_notes": recent_notes,
        "low_stock_boxes": low_stock_boxes
    })


@app.get("/cartridges", response_class=HTMLResponse)
async def cartridges_page(request: Request, db: Session = Depends(get_db)):
    cartridges = db.query(Cartridge).all()
    return templates.TemplateResponse("cartridges.html", {
        "request": request, 
        "cartridges": cartridges
    })


@app.post("/cartridges")
async def create_cartridge(
    request: Request,
    article: str = Form(...),
    model: str = Form(...),
    printer_type: str = Form(None),
    color: str = Form(None),
    status: str = Form("новый"),
    capacity: str = Form(""),
    initial_quantity: str = Form("0"),
    db: Session = Depends(get_db)
):
    qty = int(initial_quantity) if initial_quantity else 0
    cartridge = Cartridge(
        article=article,
        model=model,
        printer_type=printer_type,
        color=color,
        status=status,
        capacity=int(capacity) if capacity else None,
        initial_quantity=qty,
        total_quantity=qty
    )
    db.add(cartridge)
    db.commit()
    return RedirectResponse(url="/cartridges", status_code=303)


@app.get("/cartridges/{cartridge_id}")
async def get_cartridge(cartridge_id: int, db: Session = Depends(get_db)):
    cartridge = db.query(Cartridge).filter(Cartridge.id == cartridge_id).first()
    if not cartridge:
        raise HTTPException(status_code=404, detail="Картридж не найден")
    return {
        "id": cartridge.id,
        "article": cartridge.article,
        "model": cartridge.model,
        "printer_type": cartridge.printer_type,
        "color": cartridge.color,
        "status": cartridge.status,
        "total_quantity": cartridge.total_quantity or 0,
        "capacity": cartridge.capacity
    }


@app.post("/cartridges/{cartridge_id}/edit")
async def edit_cartridge(
    cartridge_id: int,
    article: str = Form(...),
    model: str = Form(...),
    printer_type: str = Form(""),
    color: str = Form(""),
    status: str = Form("новый"),
    capacity: str = Form(""),
    add_quantity: str = Form("0"),
    db: Session = Depends(get_db)
):
    from urllib.parse import quote
    cartridge = db.query(Cartridge).filter(Cartridge.id == cartridge_id).first()
    if not cartridge:
        return RedirectResponse(url="/cartridges?error=" + quote("Картридж не найден"), status_code=303)
    
    add_qty = int(add_quantity) if add_quantity else 0
    
    cartridge.article = article
    cartridge.model = model
    cartridge.printer_type = printer_type if printer_type else None
    cartridge.color = color if color else None
    cartridge.status = status
    cartridge.capacity = int(capacity) if capacity else None
    
    if add_qty > 0:
        cartridge.initial_quantity = (cartridge.initial_quantity or 0) + add_qty
        cartridge.total_quantity = (cartridge.total_quantity or 0) + add_qty
    
    db.commit()
    return RedirectResponse(url="/cartridges", status_code=303)


@app.post("/cartridges/{cartridge_id}/delete")
async def delete_cartridge(cartridge_id: int, db: Session = Depends(get_db)):
    from urllib.parse import quote
    cartridge = db.query(Cartridge).filter(Cartridge.id == cartridge_id).first()
    if not cartridge:
        return RedirectResponse(url="/cartridges?error=" + quote("Картридж не найден"), status_code=303)
    
    if cartridge.total_quantity and cartridge.total_quantity > 0:
        msg = f"Нельзя удалить картридж: общее количество {cartridge.total_quantity} шт. Сначала установите 0."
        return RedirectResponse(url="/cartridges?error=" + quote(msg), status_code=303)
    
    service_notes_count = db.query(ServiceNote).filter(ServiceNote.cartridge_id == cartridge_id).count()
    if service_notes_count > 0:
        msg = f"Нельзя удалить картридж: есть {service_notes_count} служебных записок. Сначала удалите их."
        return RedirectResponse(url="/cartridges?error=" + quote(msg), status_code=303)
    
    movements_count = db.query(CartridgeMovement).filter(CartridgeMovement.cartridge_id == cartridge_id).count()
    if movements_count > 0:
        db.query(CartridgeMovement).filter(CartridgeMovement.cartridge_id == cartridge_id).delete()
    
    db.delete(cartridge)
    db.commit()
    return RedirectResponse(url="/cartridges?success=" + quote("Картридж удален"), status_code=303)


@app.get("/warehouses", response_class=HTMLResponse)
async def warehouses_page(request: Request, db: Session = Depends(get_db)):
    warehouses = db.query(Warehouse).all()
    boxes = db.query(Box).all()
    cartridges = db.query(Cartridge).all()
    locations = db.query(CartridgeLocation).filter(
        CartridgeLocation.status == "на складе"
    ).all()
    
    undistributed = []
    for c in cartridges:
        distributed = sum(loc.quantity for loc in c.locations if loc.quantity)
        available = (c.total_quantity or 0) - distributed
        if available > 0:
            undistributed.append({"cartridge": c, "available": available, "distributed": distributed})
    
    return templates.TemplateResponse("warehouses.html", {
        "request": request, 
        "warehouses": warehouses,
        "boxes": boxes,
        "cartridges": cartridges,
        "locations": locations,
        "undistributed": undistributed
    })


@app.post("/warehouses")
async def create_warehouse(
    name: str = Form(...),
    location: str = Form(None),
    description: str = Form(None),
    db: Session = Depends(get_db)
):
    warehouse = Warehouse(name=name, location=location, description=description)
    db.add(warehouse)
    db.commit()
    return RedirectResponse(url="/warehouses", status_code=303)


@app.post("/boxes")
async def create_box(
    warehouse_id: int = Form(...),
    box_number: str = Form(...),
    description: str = Form(None),
    capacity: int = Form(10),
    db: Session = Depends(get_db)
):
    box = Box(
        warehouse_id=warehouse_id,
        box_number=box_number,
        description=description,
        capacity=capacity
    )
    db.add(box)
    db.commit()
    return RedirectResponse(url="/warehouses", status_code=303)


@app.get("/warehouses/{warehouse_id}")
async def get_warehouse(warehouse_id: int, db: Session = Depends(get_db)):
    warehouse = db.query(Warehouse).filter(Warehouse.id == warehouse_id).first()
    if not warehouse:
        raise HTTPException(status_code=404, detail="Склад не найден")
    return {"id": warehouse.id, "name": warehouse.name, "location": warehouse.location, "description": warehouse.description}


@app.post("/warehouses/{warehouse_id}/edit")
async def edit_warehouse(
    warehouse_id: int,
    name: str = Form(...),
    location: str = Form(None),
    description: str = Form(None),
    db: Session = Depends(get_db)
):
    warehouse = db.query(Warehouse).filter(Warehouse.id == warehouse_id).first()
    if not warehouse:
        raise HTTPException(status_code=404, detail="Склад не найден")
    warehouse.name = name
    warehouse.location = location
    warehouse.description = description
    db.commit()
    return RedirectResponse(url="/warehouses", status_code=303)


@app.post("/warehouses/{warehouse_id}/delete")
async def delete_warehouse(warehouse_id: int, db: Session = Depends(get_db)):
    from urllib.parse import quote
    warehouse = db.query(Warehouse).filter(Warehouse.id == warehouse_id).first()
    if not warehouse:
        return RedirectResponse(url="/warehouses?error=" + quote("Склад не найден"), status_code=303)
    
    boxes = db.query(Box).filter(Box.warehouse_id == warehouse_id).all()
    for box in boxes:
        notes_count = db.query(ServiceNote).filter(ServiceNote.box_id == box.id).count()
        if notes_count > 0:
            msg = f"Нельзя удалить склад: ящик {box.box_number} используется в {notes_count} служебных записках."
            return RedirectResponse(url="/warehouses?error=" + quote(msg), status_code=303)
    
    db.delete(warehouse)
    db.commit()
    return RedirectResponse(url="/warehouses?success=" + quote("Склад удален"), status_code=303)


@app.get("/boxes/{box_id}")
async def get_box(box_id: int, db: Session = Depends(get_db)):
    box = db.query(Box).filter(Box.id == box_id).first()
    if not box:
        raise HTTPException(status_code=404, detail="Ящик не найден")
    return {"id": box.id, "warehouse_id": box.warehouse_id, "box_number": box.box_number, "description": box.description, "capacity": box.capacity}


@app.post("/boxes/{box_id}/edit")
async def edit_box(
    box_id: int,
    warehouse_id: str = Form(...),
    box_number: str = Form(...),
    description: str = Form(""),
    capacity: str = Form("10"),
    db: Session = Depends(get_db)
):
    box = db.query(Box).filter(Box.id == box_id).first()
    if not box:
        raise HTTPException(status_code=404, detail="Ящик не найден")
    box.warehouse_id = int(warehouse_id) if warehouse_id else box.warehouse_id
    box.box_number = box_number
    box.description = description if description else None
    box.capacity = int(capacity) if capacity else 10
    db.commit()
    return RedirectResponse(url="/warehouses", status_code=303)


@app.post("/boxes/{box_id}/delete")
async def delete_box(box_id: int, db: Session = Depends(get_db)):
    from urllib.parse import quote
    box = db.query(Box).filter(Box.id == box_id).first()
    if not box:
        return RedirectResponse(url="/warehouses?error=" + quote("Ящик не найден"), status_code=303)
    
    notes_count = db.query(ServiceNote).filter(ServiceNote.box_id == box_id).count()
    if notes_count > 0:
        msg = f"Нельзя удалить ящик: используется в {notes_count} служебных записках. Сначала удалите их."
        return RedirectResponse(url="/warehouses?error=" + quote(msg), status_code=303)
    
    db.delete(box)
    db.commit()
    return RedirectResponse(url="/warehouses?success=" + quote("Ящик удален"), status_code=303)


@app.get("/departments", response_class=HTMLResponse)
async def departments_page(request: Request, db: Session = Depends(get_db)):
    departments = db.query(Department).all()
    return templates.TemplateResponse("departments.html", {
        "request": request, 
        "departments": departments
    })


@app.post("/departments")
async def create_department(
    name: str = Form(...),
    manager: str = Form(None),
    phone: str = Form(None),
    employee_count: int = Form(0),
    db: Session = Depends(get_db)
):
    department = Department(
        name=name, 
        manager=manager, 
        phone=phone, 
        employee_count=employee_count
    )
    db.add(department)
    db.commit()
    return RedirectResponse(url="/departments", status_code=303)


@app.get("/departments/{department_id}")
async def get_department(department_id: int, db: Session = Depends(get_db)):
    department = db.query(Department).filter(Department.id == department_id).first()
    if not department:
        raise HTTPException(status_code=404, detail="Отдел не найден")
    return {"id": department.id, "name": department.name, "manager": department.manager, "phone": department.phone, "employee_count": department.employee_count}


@app.post("/departments/{department_id}/edit")
async def edit_department(
    department_id: int,
    name: str = Form(...),
    manager: str = Form(""),
    phone: str = Form(""),
    employee_count: str = Form("0"),
    db: Session = Depends(get_db)
):
    department = db.query(Department).filter(Department.id == department_id).first()
    if not department:
        raise HTTPException(status_code=404, detail="Отдел не найден")
    department.name = name
    department.manager = manager if manager else None
    department.phone = phone if phone else None
    department.employee_count = int(employee_count) if employee_count else 0
    db.commit()
    return RedirectResponse(url="/departments", status_code=303)


@app.post("/departments/{department_id}/delete")
async def delete_department(department_id: int, db: Session = Depends(get_db)):
    from urllib.parse import quote
    department = db.query(Department).filter(Department.id == department_id).first()
    if not department:
        return RedirectResponse(url="/departments?error=" + quote("Отдел не найден"), status_code=303)
    
    employees_count = db.query(Employee).filter(Employee.department_id == department_id).count()
    if employees_count > 0:
        msg = f"Нельзя удалить отдел: в нём {employees_count} сотрудник(ов). Сначала переведите их в другой отдел."
        return RedirectResponse(url="/departments?error=" + quote(msg), status_code=303)
    
    db.delete(department)
    db.commit()
    return RedirectResponse(url="/departments?success=" + quote("Отдел удален"), status_code=303)


@app.get("/employees", response_class=HTMLResponse)
async def employees_page(request: Request, db: Session = Depends(get_db)):
    employees = db.query(Employee).all()
    departments = db.query(Department).all()
    return templates.TemplateResponse("employees.html", {
        "request": request, 
        "employees": employees,
        "departments": departments
    })


@app.post("/employees")
async def create_employee(
    full_name: str = Form(...),
    position: str = Form(None),
    department_id: int = Form(None),
    personnel_number: str = Form(None),
    phone: str = Form(None),
    email: str = Form(None),
    db: Session = Depends(get_db)
):
    employee = Employee(
        full_name=full_name,
        position=position,
        department_id=department_id,
        personnel_number=personnel_number,
        phone=phone,
        email=email
    )
    db.add(employee)
    db.commit()
    return RedirectResponse(url="/employees", status_code=303)


@app.get("/employees/{employee_id}")
async def get_employee(employee_id: int, db: Session = Depends(get_db)):
    employee = db.query(Employee).filter(Employee.id == employee_id).first()
    if not employee:
        raise HTTPException(status_code=404, detail="Сотрудник не найден")
    return {"id": employee.id, "full_name": employee.full_name, "position": employee.position, "department_id": employee.department_id, "personnel_number": employee.personnel_number, "phone": employee.phone, "email": employee.email}


@app.post("/employees/{employee_id}/edit")
async def edit_employee(
    employee_id: int,
    full_name: str = Form(...),
    position: str = Form(""),
    department_id: str = Form(""),
    personnel_number: str = Form(""),
    phone: str = Form(""),
    email: str = Form(""),
    db: Session = Depends(get_db)
):
    employee = db.query(Employee).filter(Employee.id == employee_id).first()
    if not employee:
        raise HTTPException(status_code=404, detail="Сотрудник не найден")
    employee.full_name = full_name
    employee.position = position if position else None
    employee.department_id = int(department_id) if department_id else None
    employee.personnel_number = personnel_number if personnel_number else None
    employee.phone = phone if phone else None
    employee.email = email if email else None
    db.commit()
    return RedirectResponse(url="/employees", status_code=303)


@app.post("/employees/{employee_id}/delete")
async def delete_employee(employee_id: int, db: Session = Depends(get_db)):
    from urllib.parse import quote
    employee = db.query(Employee).filter(Employee.id == employee_id).first()
    if not employee:
        return RedirectResponse(url="/employees?error=" + quote("Сотрудник не найден"), status_code=303)
    
    notes_count = db.query(ServiceNote).filter(
        (ServiceNote.author_id == employee_id) | (ServiceNote.recipient_id == employee_id)
    ).count()
    if notes_count > 0:
        msg = f"Нельзя удалить сотрудника: участвует в {notes_count} служебных записках. Сначала удалите их."
        return RedirectResponse(url="/employees?error=" + quote(msg), status_code=303)
    
    db.query(CartridgeLocation).filter(CartridgeLocation.employee_id == employee_id).update({"employee_id": None})
    
    db.delete(employee)
    db.commit()
    return RedirectResponse(url="/employees?success=" + quote("Сотрудник удален"), status_code=303)


@app.get("/service-notes", response_class=HTMLResponse)
async def service_notes_page(request: Request, db: Session = Depends(get_db)):
    notes = db.query(ServiceNote).order_by(ServiceNote.created_date.desc()).all()
    employees = db.query(Employee).all()
    departments = db.query(Department).all()
    cartridges = db.query(Cartridge).all()
    boxes = db.query(Box).filter(Box.current_count > 0).all()
    locations = db.query(CartridgeLocation).filter(
        CartridgeLocation.status == "на складе",
        CartridgeLocation.quantity > 0
    ).all()
    
    return templates.TemplateResponse("service_notes.html", {
        "request": request, 
        "notes": notes,
        "employees": employees,
        "departments": departments,
        "cartridges": cartridges,
        "boxes": boxes,
        "locations": locations
    })


@app.post("/service-notes")
async def create_service_note(
    recipient_id: int = Form(...),
    cartridge_id: int = Form(...),
    quantity: int = Form(1),
    reason: str = Form(...),
    comment: str = Form(None),
    db: Session = Depends(get_db)
):
    from urllib.parse import quote
    
    location = db.query(CartridgeLocation).filter(
        CartridgeLocation.cartridge_id == cartridge_id,
        CartridgeLocation.status == "на складе",
        CartridgeLocation.quantity >= quantity
    ).order_by(CartridgeLocation.quantity.desc()).first()
    
    if not location:
        return RedirectResponse(url="/service-notes?error=" + quote("Недостаточно картриджей на складе"), status_code=303)
    
    cartridge = db.query(Cartridge).filter(Cartridge.id == cartridge_id).first()
    if not cartridge or (cartridge.total_quantity or 0) < quantity:
        return RedirectResponse(url="/service-notes?error=" + quote("Недостаточно картриджей"), status_code=303)
    
    note_number = generate_note_number(db)
    box_id = location.box_id
    
    note = ServiceNote(
        note_number=note_number,
        author_id=None,
        recipient_id=recipient_id,
        cartridge_id=cartridge_id,
        quantity=quantity,
        box_id=box_id,
        reason=reason,
        comment=comment,
        status="выдано"
    )
    db.add(note)
    
    location.quantity -= quantity
    if location.quantity <= 0:
        location.status = "выдано"
    
    cartridge.total_quantity = max(0, (cartridge.total_quantity or 0) - quantity)
    
    box = db.query(Box).filter(Box.id == box_id).first()
    if box:
        box.current_count = max(0, box.current_count - quantity)
    
    recipient = db.query(Employee).filter(Employee.id == recipient_id).first()
    
    movement = CartridgeMovement(
        cartridge_id=cartridge_id,
        from_location=f"Ящик {box.box_number}" if box else "Склад",
        to_location=f"Сотрудник: {recipient.full_name}" if recipient else "Выдано",
        service_note_id=note.id
    )
    db.add(movement)
    
    db.commit()
    return RedirectResponse(url="/service-notes?success=" + quote(f"Служебка {note_number} создана"), status_code=303)


@app.post("/service-notes/{note_id}/return")
async def return_cartridge(note_id: int, db: Session = Depends(get_db)):
    note = db.query(ServiceNote).filter(ServiceNote.id == note_id).first()
    if not note:
        raise HTTPException(status_code=404, detail="Служебка не найдена")
    
    note.status = "возврат"
    
    location = db.query(CartridgeLocation).filter(
        CartridgeLocation.cartridge_id == note.cartridge_id,
        CartridgeLocation.box_id == note.box_id
    ).first()
    
    if location:
        location.quantity += note.quantity
        location.status = "на складе"
    else:
        new_location = CartridgeLocation(
            cartridge_id=note.cartridge_id,
            box_id=note.box_id,
            status="на складе",
            quantity=note.quantity
        )
        db.add(new_location)
    
    box = db.query(Box).filter(Box.id == note.box_id).first()
    if box:
        box.current_count += note.quantity
    
    movement = CartridgeMovement(
        cartridge_id=note.cartridge_id,
        from_location="В использовании",
        to_location=f"Ящик {box.box_number}" if box else "Склад",
        service_note_id=note.id
    )
    db.add(movement)
    
    db.commit()
    return RedirectResponse(url="/service-notes", status_code=303)


@app.post("/locations")
async def add_cartridge_to_stock(
    cartridge_id: int = Form(...),
    box_id: int = Form(...),
    quantity: int = Form(1),
    db: Session = Depends(get_db)
):
    from urllib.parse import quote
    
    cartridge = db.query(Cartridge).filter(Cartridge.id == cartridge_id).first()
    if not cartridge:
        return RedirectResponse(url="/warehouses?error=" + quote("Картридж не найден"), status_code=303)
    
    current_distributed = sum(loc.quantity for loc in cartridge.locations if loc.quantity)
    available = (cartridge.total_quantity or 0) - current_distributed
    
    if quantity > available:
        msg = f"Нельзя добавить {quantity} шт.: доступно только {available} шт. (всего {cartridge.total_quantity}, распределено {current_distributed})"
        return RedirectResponse(url="/warehouses?error=" + quote(msg), status_code=303)
    
    box = db.query(Box).filter(Box.id == box_id).first()
    if not box:
        return RedirectResponse(url="/warehouses?error=" + quote("Ящик не найден"), status_code=303)
    
    free_space = box.capacity - (box.current_count or 0)
    if quantity > free_space:
        msg = f"Нельзя добавить {quantity} шт. в ящик {box.box_number}: свободно только {free_space} мест (вместимость {box.capacity}, занято {box.current_count})"
        return RedirectResponse(url="/warehouses?error=" + quote(msg), status_code=303)
    
    location = db.query(CartridgeLocation).filter(
        CartridgeLocation.cartridge_id == cartridge_id,
        CartridgeLocation.box_id == box_id,
        CartridgeLocation.status == "на складе"
    ).first()
    
    if location:
        location.quantity += quantity
    else:
        location = CartridgeLocation(
            cartridge_id=cartridge_id,
            box_id=box_id,
            status="на складе",
            quantity=quantity
        )
        db.add(location)
    
    box.current_count += quantity
    
    movement = CartridgeMovement(
        cartridge_id=cartridge_id,
        from_location="Поступление",
        to_location=f"Ящик {box.box_number}" if box else "Склад"
    )
    db.add(movement)
    
    db.commit()
    return RedirectResponse(url="/warehouses?success=" + quote(f"Добавлено {quantity} шт."), status_code=303)


@app.post("/locations/{location_id}/remove-one")
async def remove_one_from_location(location_id: int, db: Session = Depends(get_db)):
    from urllib.parse import quote
    
    location = db.query(CartridgeLocation).filter(CartridgeLocation.id == location_id).first()
    if not location:
        return RedirectResponse(url="/warehouses?error=" + quote("Запись не найдена"), status_code=303)
    
    box = db.query(Box).filter(Box.id == location.box_id).first()
    cartridge = location.cartridge
    
    if location.quantity <= 1:
        db.delete(location)
    else:
        location.quantity -= 1
    
    if box and box.current_count > 0:
        box.current_count -= 1
    
    movement = CartridgeMovement(
        cartridge_id=location.cartridge_id,
        from_location=f"Ящик {box.box_number}" if box else "Склад",
        to_location="Нераспределённые"
    )
    db.add(movement)
    
    db.commit()
    return RedirectResponse(url="/warehouses?success=" + quote(f"Убран 1 шт. {cartridge.article} в нераспределённые"), status_code=303)


@app.get("/reports", response_class=HTMLResponse)
async def reports_page(request: Request, db: Session = Depends(get_db)):
    inventory = db.query(
        Box.box_number,
        Warehouse.name.label("warehouse_name"),
        func.sum(CartridgeLocation.quantity).label("total")
    ).join(Warehouse).outerjoin(CartridgeLocation).filter(
        CartridgeLocation.status == "на складе"
    ).group_by(Box.id, Warehouse.name).all()
    
    dept_stats = db.query(
        Department.name,
        func.count(ServiceNote.id).label("notes_count")
    ).join(Employee, Employee.department_id == Department.id).join(
        ServiceNote, ServiceNote.recipient_id == Employee.id
    ).group_by(Department.id).all()
    
    low_stock = db.query(Box).filter(Box.current_count < 3).all()
    
    movements = db.query(CartridgeMovement).order_by(
        CartridgeMovement.movement_date.desc()
    ).limit(20).all()
    
    return templates.TemplateResponse("reports.html", {
        "request": request,
        "inventory": inventory,
        "dept_stats": dept_stats,
        "low_stock": low_stock,
        "movements": movements
    })


@app.get("/api/employee/{employee_id}/department")
async def get_employee_department(employee_id: int, db: Session = Depends(get_db)):
    employee = db.query(Employee).filter(Employee.id == employee_id).first()
    if employee and employee.department:
        return {"department_name": employee.department.name}
    return {"department_name": ""}


@app.get("/service-notes/{note_id}/print")
async def print_service_note(note_id: int, db: Session = Depends(get_db)):
    note = db.query(ServiceNote).filter(ServiceNote.id == note_id).first()
    if not note:
        raise HTTPException(status_code=404, detail="Служебка не найдена")
    
    doc = Document()
    
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)
    
    author_position = note.author.position if note.author and note.author.position else ""
    if note.author and note.author.full_name:
        name_parts = note.author.full_name.split()
        if len(name_parts) >= 3:
            author_short = f"{name_parts[0]} {name_parts[1][0]}. {name_parts[2][0]}."
        elif len(name_parts) == 2:
            author_short = f"{name_parts[0]} {name_parts[1][0]}."
        else:
            author_short = note.author.full_name
    else:
        author_short = "_______________"
    
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.left_indent = Cm(8.25)
    header.paragraph_format.right_indent = Cm(-1.76)
    header.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    header.add_run("Исполняющему обязанности начальника\nотдела информатизации\nКраснодарского филиала\nРЭУ им. Г.В. Плеханова\nПетровскому Д. А.\n")
    if author_position:
        header.add_run(f"{author_position}\n")
    header.add_run(f"{author_short}")
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("служебная записка.")
    run.font.size = Pt(14)
    
    doc.add_paragraph()
    
    cartridge_info = f"{note.cartridge.article}" if note.cartridge else "картридж"
    cartridge_model = f"{note.cartridge.model}" if note.cartridge else ""
    printer_type = note.cartridge.printer_type if note.cartridge and note.cartridge.printer_type else "лазерного"
    
    body = doc.add_paragraph()
    body.alignment = WD_ALIGN_PARAGRAPH.CENTER
    body.add_run(f"Прошу предоставить картридж {cartridge_info} для лазерного принтера {cartridge_model} в количестве {note.quantity} шт.")
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    signature = doc.add_paragraph()
    signature.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    signature.add_run("_______________________")
    
    temp_dir = tempfile.gettempdir()
    filename = f"service_note_{note.note_number.replace('-', '_')}.docx"
    filepath = os.path.join(temp_dir, filename)
    doc.save(filepath)
    
    return FileResponse(
        path=filepath,
        filename=filename,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=5000)
